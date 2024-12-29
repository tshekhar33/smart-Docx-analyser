import streamlit as st
import markdown
import pdfkit
import os
import PyPDF2
from pathlib import Path
import tempfile
import logging
import re
from bs4 import BeautifulSoup
from textblob import TextBlob
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
import nltk
from nltk.tokenize import sent_tokenize
import base64
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import mammoth

# Initialize theme in session state if not present
if 'theme' not in st.session_state:
    st.session_state.theme = 'light'

# Custom CSS with dynamic theme
def get_theme_css():
    light_theme = {
        'bg_color': '#fce4ec',
        'text_color': '#000000',
        'heading_color': '#000000',
        'title_color': '#000000',
        'subheading_color': '#000000',
        'button_bg': '#1976d2',
        'button_text': '#ffffff',
        'button_hover': '#1565c0',
        'sidebar_bg': '#f8bbd0',
        'card_bg': '#ffffff',
        'accent_color': '#000000',
        'success_color': '#000000',
        'error_color': '#000000',
        'link_color': '#000000',
        'upload_bg': '#ffffff'
    }
    
    dark_theme = {
        'bg_color': '#1a1a1a',
        'text_color': '#ffffff',
        'heading_color': '#ffffff',
        'title_color': '#ffffff',
        'subheading_color': '#ffffff',
        'button_bg': '#1976d2',
        'button_text': '#ffffff',
        'button_hover': '#1565c0',
        'sidebar_bg': '#212121',
        'card_bg': '#2d2d2d',
        'accent_color': '#ffffff',
        'success_color': '#ffffff',
        'error_color': '#ffffff',
        'link_color': '#ffffff',
        'upload_bg': '#ffffff'
    }
    
    theme = dark_theme if st.session_state.theme == 'dark' else light_theme
    
    return f"""
    <style>
        /* Main theme */
        .main {{
            background-color: {theme['bg_color']};
            color: {theme['text_color']};
        }}
        .stApp {{
            background-color: {theme['bg_color']};
        }}
        
        /* App title */
        .main h1:first-of-type {{
            color: {theme['title_color']} !important;
            font-size: 2.5rem !important;
            font-weight: bold !important;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
            padding: 1rem;
            background-color: {theme['card_bg']};
            border-radius: 10px;
            margin-bottom: 2rem;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        
        /* Text colors */
        p, li, .stMarkdown, label, span {{
            color: {theme['text_color']} !important;
        }}
        
        /* Headings */
        h1, h2, h3, h4, h5 {{
            color: {theme['heading_color']} !important;
            font-weight: bold;
        }}
        
        /* File uploader styling */
        .stFileUploader {{
            background-color: {theme['upload_bg']} !important;
        }}
        
        .stFileUploader > div {{
            background-color: {theme['upload_bg']} !important;
            color: #000000 !important;
        }}
        
        .stFileUploader > div > div {{
            background-color: {theme['upload_bg']} !important;
            color: #000000 !important;
        }}
        
        /* Browse files button */
        .stFileUploader button {{
            background-color: {theme['button_bg']} !important;
            color: {theme['button_text']} !important;
            border: none !important;
            padding: 0.5rem 1rem !important;
            border-radius: 5px !important;
            font-weight: bold !important;
            cursor: pointer !important;
            transition: all 0.3s ease !important;
        }}
        
        .stFileUploader button:hover {{
            background-color: {theme['button_hover']} !important;
            transform: translateY(-2px) !important;
            box-shadow: 0 4px 8px rgba(0,0,0,0.2) !important;
        }}
        
        /* Download button */
        .stDownloadButton button {{
            background-color: {theme['button_bg']} !important;
            color: {theme['button_text']} !important;
            border: none !important;
            padding: 0.75rem 1.5rem !important;
            border-radius: 5px !important;
            font-weight: bold !important;
            cursor: pointer !important;
            transition: all 0.3s ease !important;
            width: 100% !important;
            margin: 4px 2px !important;
        }}
        
        .stDownloadButton button:hover {{
            background-color: {theme['button_hover']} !important;
            transform: translateY(-2px) !important;
            box-shadow: 0 4px 8px rgba(0,0,0,0.2) !important;
        }}
        
        /* Convert and Analyze buttons */
        .stButton > button {{
            background-color: {theme['button_bg']} !important;
            color: {theme['button_text']} !important;
            border: none !important;
            padding: 0.75rem 1.5rem !important;
            border-radius: 5px !important;
            font-weight: bold !important;
            cursor: pointer !important;
            transition: all 0.3s ease !important;
            width: 100% !important;
        }}
        
        .stButton > button:hover {{
            background-color: {theme['button_hover']} !important;
            transform: translateY(-2px) !important;
            box-shadow: 0 4px 8px rgba(0,0,0,0.2) !important;
        }}
        
        /* Upload text color */
        .uploadedFileName {{
            color: #000000 !important;
        }}
        
        /* File uploader text */
        .stFileUploader p {{
            color: #000000 !important;
        }}
        
        /* Buttons */
        .stButton>button {{
            background-color: {theme['button_bg']};
            color: {theme['button_text']};
            border-radius: 25px;
            padding: 0.5rem 1.5rem;
            border: none;
            transition: all 0.3s;
            font-weight: bold;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        .stButton>button:hover {{
            background-color: {theme['button_hover']};
            transform: translateY(-2px);
            box-shadow: 0 4px 8px rgba(0,0,0,0.2);
        }}
        
        /* Sidebar */
        [data-testid="stSidebar"] {{
            background-color: {theme['sidebar_bg']};
        }}
        [data-testid="stSidebar"] h1,
        [data-testid="stSidebar"] h2,
        [data-testid="stSidebar"] h3,
        [data-testid="stSidebar"] h4,
        [data-testid="stSidebar"] p,
        [data-testid="stSidebar"] li {{
            color: {theme['text_color']} !important;
        }}
        
        /* Cards and containers */
        .stTextArea textarea,
        .uploadedFile,
        .streamlit-expanderContent {{
            background-color: {theme['card_bg']} !important;
            color: {theme['text_color']} !important;
            border: 1px solid #ddd;
        }}
        
        /* Metrics */
        [data-testid="stMetricValue"] {{
            color: {theme['text_color']} !important;
            font-weight: bold;
        }}
        [data-testid="stMetricLabel"] {{
            color: {theme['text_color']} !important;
        }}
        
        /* File uploader */
        .stFileUploader {{
            background-color: {theme['card_bg']};
            border: 2px dashed #ddd;
            color: {theme['text_color']} !important;
        }}
        
        /* Expander */
        .streamlit-expanderHeader {{
            color: {theme['text_color']} !important;
            font-weight: bold;
        }}
        
        /* Links */
        a {{
            color: {theme['text_color']} !important;
            text-decoration: underline;
        }}
        
        /* Tabs */
        .stTabs [data-baseweb="tab"] {{
            color: {theme['text_color']};
        }}
        .stTabs [aria-selected="true"] {{
            background-color: {theme['button_bg']};
            color: {theme['button_text']} !important;
        }}
        
        /* Upload text */
        .stUploadMessage p {{
            color: {theme['text_color']} !important;
        }}
        
        /* Progress bar text */
        .stProgress p {{
            color: {theme['text_color']} !important;
        }}
    </style>
    """

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set page configuration
st.set_page_config(
    page_title="Smart Document Processor",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

def get_text_summary(text, sentences_count=3):
    """Generate a summary of the text using LSA."""
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, sentences_count)
    return " ".join([str(sentence) for sentence in summary])

def analyze_sentiment(text):
    """Analyze the sentiment of the text."""
    blob = TextBlob(text)
    sentiment = blob.sentiment.polarity
    if sentiment > 0:
        return "Positive", sentiment
    elif sentiment < 0:
        return "Negative", sentiment
    else:
        return "Neutral", sentiment

def extract_key_phrases(text):
    """Extract important phrases from the text."""
    blob = TextBlob(text)
    phrases = blob.noun_phrases
    return list(set(phrases))

def analyze_content_structure(md_content):
    """Analyze the content structure and identify different sections."""
    chapters = re.findall(r'CHAPTER[- ]\w+', md_content)
    sections = re.findall(r'\d{3}[- ][A-Z][^\n]+', md_content)
    subsections = re.findall(r'\d{3}\.\d+[- ][A-Z][^\n]+', md_content)
    
    return {
        'chapters': chapters,
        'sections': sections,
        'subsections': subsections
    }

def format_content(md_content):
    """Format the content based on its structure."""
    # Format chapters
    md_content = re.sub(
        r'(CHAPTER[- ]\w+)',
        r'<div class="chapter">\n# \1\n</div>',
        md_content
    )
    
    # Format sections
    md_content = re.sub(
        r'(\d{3})[- ]([A-Z][^\n]+)',
        r'## \1 \2',
        md_content
    )
    
    # Format subsections
    md_content = re.sub(
        r'(\d{3}\.\d+)[- ]([A-Z][^\n]+)',
        r'### \1 \2',
        md_content
    )
    
    # Format lists
    md_content = re.sub(r'(?m)^\s*(\d+\)|\w\))\s*', r'1. ', md_content)
    
    # Add spacing
    md_content = re.sub(r'\n\n', r'\n\n<br>\n\n', md_content)
    
    return md_content

# Custom button styles
download_button_css = """
<style>
.download-button {
    background-color: #1976d2;
    color: black !important;
    padding: 0.5rem 1rem;
    border-radius: 5px;
    border: none;
    font-weight: bold;
    cursor: pointer;
    width: 100%;
    text-align: center;
    text-decoration: none;
    display: inline-block;
    margin: 4px 2px;
    transition: all 0.3s ease;
}
.download-button:hover {
    background-color: #1565c0;
    transform: translateY(-2px);
    box-shadow: 0 4px 8px rgba(0,0,0,0.2);
}
</style>
"""

def setup_document_styles(doc):
    """Set up document styles for professional formatting"""
    styles = doc.styles

    # Heading 1 style
    h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    h1_format = h1_style.paragraph_format
    h1_format.space_before = Pt(24)
    h1_format.space_after = Pt(12)
    h1_format.keep_with_next = True
    h1_font = h1_style.font
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0)

    # Heading 2 style
    h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_format = h2_style.paragraph_format
    h2_format.space_before = Pt(18)
    h2_format.space_after = Pt(6)
    h2_format.keep_with_next = True
    h2_font = h2_style.font
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)

    # Normal text style
    normal_style = styles.add_style('Custom Normal', WD_STYLE_TYPE.PARAGRAPH)
    normal_format = normal_style.paragraph_format
    normal_format.space_after = Pt(12)
    normal_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_font = normal_style.font
    normal_font.size = Pt(12)
    normal_font.name = 'Times New Roman'

def convert_md_to_pdf_enhanced(md_content, output_filename):
    """Enhanced conversion from Markdown to PDF with better formatting"""
    try:
        # First, create a DOCX file
        doc = Document()
        setup_document_styles(doc)

        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Convert markdown to HTML first
        html_content = markdown.markdown(md_content, extensions=['tables', 'toc'])
        soup = BeautifulSoup(html_content, 'html.parser')

        # Process the content
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol']):
            if element.name == 'h1':
                p = doc.add_paragraph(element.get_text(), style='Custom Heading 1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif element.name == 'h2':
                p = doc.add_paragraph(element.get_text(), style='Custom Heading 2')
            elif element.name == 'h3':
                p = doc.add_paragraph(element.get_text(), style='Custom Heading 2')
            elif element.name == 'p':
                p = doc.add_paragraph(element.get_text(), style='Custom Normal')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    p = doc.add_paragraph(style='Custom Normal')
                    p.add_run('• ' + li.get_text())
                    p.paragraph_format.left_indent = Inches(0.25)

        # Save as DOCX first
        temp_docx = "temp_output.docx"
        doc.save(temp_docx)

        # Convert DOCX to PDF
        config = pdfkit.configuration(wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')
        
        # Convert DOCX to HTML using mammoth
        with open(temp_docx, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value

        # Add CSS for better PDF styling
        html_with_style = f"""
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2.5cm 2cm;
                }}
                body {{
                    font-family: "Times New Roman", Times, serif;
                    font-size: 12pt;
                    line-height: 1.5;
                }}
                h1 {{
                    font-size: 16pt;
                    text-align: center;
                    margin-top: 24pt;
                    margin-bottom: 12pt;
                }}
                h2 {{
                    font-size: 14pt;
                    margin-top: 18pt;
                    margin-bottom: 6pt;
                }}
                p {{
                    text-align: justify;
                    margin-bottom: 12pt;
                }}
                ul, ol {{
                    margin-left: 0.25in;
                }}
            </style>
        </head>
        <body>
            {html}
        </body>
        </html>
        """

        # Convert to PDF
        pdfkit.from_string(html_with_style, output_filename, configuration=config)
        
        # Clean up temporary file
        os.remove(temp_docx)
        return True
    except Exception as e:
        logging.error(f"Error in enhanced conversion: {str(e)}")
        return False

def main():
    # Initialize theme if not present
    if 'theme' not in st.session_state:
        st.session_state.theme = 'light'
        
    # Apply theme CSS
    st.markdown(get_theme_css(), unsafe_allow_html=True)
    
    st.title("📄 Smart Document Processor")
    
    # Theme switcher in sidebar
    with st.sidebar:
        st.markdown("## Settings")
        theme_switch = st.toggle("Dark Mode", value=st.session_state.theme == 'dark', key='theme_toggle')
        if theme_switch != (st.session_state.theme == 'dark'):
            st.session_state.theme = 'dark' if theme_switch else 'light'
            st.rerun()
        
        st.markdown("## About", unsafe_allow_html=True)
        st.markdown("""
        This application helps you process documents intelligently:
        * Convert Markdown to PDF
        * Analyze document structure
        * Generate summaries
        * Extract key information
        """)
        
        st.markdown("## Statistics", unsafe_allow_html=True)
        st.markdown(f"**Current time:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Main content
    tabs = st.tabs(["📝 Convert & Analyze", "🔍 Document Analysis", "ℹ️ Help"])
    
    with tabs[0]:
        st.header("Document Conversion")
        
        # Create a container for file upload and buttons
        with st.container():
            col1, col2 = st.columns([3, 1])
            
            with col1:
                uploaded_file = st.file_uploader(
                    "Upload your document",
                    type=['md', 'txt'],
                    help="Upload a Markdown or text file"
                )
            
            if uploaded_file:
                content = uploaded_file.getvalue().decode()
                
                # Show document structure with colored metrics
                with st.expander("📊 Document Structure", expanded=True):
                    structure = analyze_content_structure(content)
                    cols = st.columns(3)
                    with cols[0]:
                        st.metric("📚 Chapters", len(structure['chapters']))
                    with cols[1]:
                        st.metric("📑 Sections", len(structure['sections']))
                    with cols[2]:
                        st.metric("📎 Subsections", len(structure['subsections']))
                
                # Action buttons
                col3, col4, col5 = st.columns([1, 1, 1])
                with col3:
                    convert_button = st.button("🔄 Convert to PDF", use_container_width=True, type="primary")
                with col4:
                    analyze_button = st.button("🔍 Analyze", use_container_width=True, type="secondary")
                with col5:
                    preview_button = st.button("👁️ Preview", use_container_width=True)
                
                # Content Preview
                if preview_button:
                    st.markdown("### 📄 Content Preview")
                    st.text_area(
                        "Document Content",
                        value=content,
                        height=300,
                        key="preview"
                    )
                
                # AI Analysis with colored sections
                if analyze_button:
                    with st.expander("🤖 AI Analysis", expanded=True):
                        col_a, col_b = st.columns(2)
                        
                        with col_a:
                            st.markdown("### 📝 Content Summary")
                            summary = get_text_summary(content)
                            st.write(summary)
                        
                        with col_b:
                            st.markdown("### 🎯 Key Metrics")
                            # Sentiment analysis
                            sentiment, score = analyze_sentiment(content)
                            st.markdown(f"**Sentiment:** {sentiment}")
                            st.progress(float(score + 1) / 2, text=f"Score: {score:.2f}")
                            
                            # Key phrases
                            st.markdown("### 🔑 Key Phrases")
                            phrases = extract_key_phrases(content)
                            st.markdown(", ".join(f"**{phrase}**" for phrase in phrases[:5]))
                
                # Convert to PDF
                if convert_button:
                    with st.spinner("Converting..."):
                        output_filename = "output.pdf"
                        if convert_md_to_pdf_enhanced(content, output_filename):
                            with open(output_filename, "rb") as pdf_file:
                                pdf_bytes = pdf_file.read()
                                st.success("✅ Conversion successful!")
                                
                                # Add custom styling for download button
                                st.markdown(download_button_css, unsafe_allow_html=True)
                                
                                # Create custom download button
                                col1, col2, col3 = st.columns([1, 2, 1])
                                with col2:
                                    st.download_button(
                                        label="📥 Download PDF",
                                        data=pdf_bytes,
                                        file_name=output_filename,
                                        mime="application/pdf",
                                        use_container_width=True,
                                        key="download_button"
                                    )
                            os.remove(output_filename)
                        else:
                            st.error("❌ Conversion failed")
    
    with tabs[1]:
        st.header("PDF Analysis")
        pdf_file = st.file_uploader("Upload PDF for analysis", type=['pdf'])
        
        if pdf_file:
            with st.spinner("Analyzing PDF..."):
                reader = PyPDF2.PdfReader(pdf_file)
                num_pages = len(reader.pages)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Total Pages", num_pages)
                with col2:
                    file_size_mb = len(pdf_file.getvalue()) / (1024 * 1024)
                    st.metric("File Size", f"{file_size_mb:.1f} MB")
                
                # Extract and analyze text
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                
                with st.expander("📊 Content Analysis"):
                    # Summary
                    st.subheader("Summary")
                    st.write(get_text_summary(text))
                    
                    # Sentiment
                    sentiment, score = analyze_sentiment(text)
                    st.subheader("Document Sentiment")
                    st.progress(float(score + 1) / 2, text=f"Score: {score:.2f}")
                    st.write(f"Overall tone: {sentiment}")
                    
                    # Key phrases
                    st.subheader("Key Topics")
                    phrases = extract_key_phrases(text)
                    st.write(", ".join(f"**{phrase}**" for phrase in phrases[:5]))
    
    with tabs[2]:
        st.header("Help & Documentation")
        st.markdown("""
        ### How to Use
        1. **Upload** your Markdown or text document
        2. Review the **AI analysis** and structure
        3. Click **Convert to PDF** to generate a formatted PDF
        4. **Download** the converted file
        
        ### Features
        - 📊 Document structure analysis
        - 🤖 AI-powered content analysis
        - 📝 Smart formatting
        - 📄 Professional PDF conversion
        
        ### Need Help?
        Contact support at support@example.com
        """)

if __name__ == "__main__":
    main()
